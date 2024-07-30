# 1.
# read from a '.txt' file and count the number of sentences, words, and characters in the file.
# count no. of sentences
def countLines(file):
    f = open(file, 'r')
    dat = f.read().split('.')
    f.close()
    dat.pop()
    count = len(dat)
    return count
# count no. of words
def countWords(file):
    f = open(file, 'r')
    count = 0
    line = f.readline()
    while line:
        count += len(line.split())
        line = f.readline()
    f.close()
    return count
# count no. of characters
def countChars(file):
    f = open(file, 'r')
    count = 0
    line = f.readline()
    while line:
        for word in line.split():
            for i in word:
                if i.isalnum():
                    count += 1
        line = f.readline()
    f.close()
    return count
print('No. of Sentences = ', countLines(r'File1.txt'))
print('No. of Words = ', countWords(r'File1.txt'))
print('No. of Characters = ', countChars(r'File1.txt'))

# 2.
# read from a '.pdf' file and count the number of sentences, words, and characters in the file.
from pypdf import PdfReader
# count no. of sentences
def countLines(file):
    f = PdfReader(file)
    count = 0
    for page in f.pages:
        dat = page.extract_text().split('.')
        dat.pop()
        count += len(dat)
    return count
# count no. of words
def countWords(file):
    f = PdfReader(file)
    count = 0
    for page in f.pages:
        dat = page.extract_text().split(' ')
        for i in dat:
            if i != '':
                count += 1
    return count
# count no. of characters
def countChars(file):
    f = PdfReader(file)
    count = 0
    for page in f.pages:
        dat = page.extract_text()
        for i in dat:
            if i.isalnum():
                count += 1
    return count
print('No. of Sentences = ', countLines(r'example.pdf'))
print('No. of Words = ', countWords(r'example.pdf'))
print('No. of Characters = ', countChars(r'example.pdf'))
#3.
# read from a '.docx' file and count the number of sentences, words, and characters in the file.
from docx import Document
# count no. of sentences
def countLines(file):
    f = Document(file)
    count = 0
    for para in f.paragraphs:
        dat = para.text.split('.')
        dat.pop()
        count += len(dat)
    return count
# count no. of words
def countWords(file):
    f = Document(file)
    count = 0
    for para in f.paragraphs:
        dat = para.text.split(' ')
        for i in dat:
            if i != '':
                count += 1
    return count
# count no. of characters
def countChars(file):
    f = Document(file)
    count = 0
    for para in f.paragraphs:
        dat = para.text
        for i in dat:
            if i.isalnum():
                count += 1
    return count
print('No. of Sentences = ', countLines(r'DSNotesCh-4.docx'))
print('No. of Words = ', countWords(r'DSNotesCh-4.docx'))
print('No. of Characters = ', countChars(r'DSNotesCh-4.docx'))

#4.
# Implement term-document incidence matrix for boolean retrieval.
import os
def getFiles():
    lst = []
    for f in os.listdir():      # gives list of files in current directory
        if f.endswith('.txt'):  # taking only .txt files from current directory
            lst.append(f)
    return lst
# create term-document matrix
def createMatrix(files):
    fp = []     # index of this list work as documentID
    for file in files:
        f = open(file, 'r')
        fp.append(f)
    words = []  # to store unique words which acts as rows
    matrix = []
    for i, f in enumerate(fp):
        dat = f.read().split()
        for word in dat:
            if word not in words:
                words.append(word)
                matrix.append([0 for _ in range(len(fp))])
            matrix[words.index(word)][i] = 1
    return matrix, words
# display term-document matrix
def printMatrix(matrix, words, files):
    print('\t', end='\t')
    for i in files:
        print(i, end='\t')
    print()
    for i in range(len(words)):
        print(words[i], ' '*(16-len(words[i])) ,end='')
        for j in range(len(files)):
            print(matrix[i][j], end='\t\t')
        print()
# perform AND Query
def andQuery(w1, w2, matrix, words):
    try:    # gets row of word-1 if available, else error
        l1 = matrix[words.index(w1)]
        l1 = int(''.join(map(str, l1)), 2)
    except ValueError:
        return -1
    try:    # gets row of word-2 if available, else error
        l2 = matrix[words.index(w2)]
        l2 = int(''.join(map(str, l2)), 2)
    except ValueError:
        return -2
    return bin(l1 & l2)[2:]     # bitwise and
# perform OR Query
def orQuery(w1, w2, matrix, words):
    try:    # gets row of word-1 if available, else works with word-2 only
        l1 = matrix[words.index(w1)]
        l1 = int(''.join(map(str, l1)), 2)
    except ValueError:
        print(f'{w1} Not Found')
        l1 = 0
    try:    # gets row of word-2 if available, else works with word-1 only
        l2 = matrix[words.index(w2)]
        l2 = int(''.join(map(str, l2)), 2)
    except ValueError:
        print(f'{w2} Not Found')
        l2 = 0
    return bin(l1 | l2)[2:]     # bitwise or
# perform NOT Query
def notQuery(w, matrix, words):
    try:    # gets row of word if available, else gives all 0's
        l = matrix[words.index(w)]
        l = ''.join(map(str, l))
    except:
        l = '0' * len(matrix[0])
    return ''.join(['0' if i=='1' else '1' for i in l])     # complementing
if __name__ == '__main__':
    files = getFiles()
    matrix, words = createMatrix(files)
    print('Term-Document Incidence Matrix: ')
    printMatrix(matrix, words, files)
    print('Vocabulary:')
    print(words)
    print('Terms in lowercase & Operator in uppercase')
    query = input('Enter query: ').split()
    if len(query) == 3:
        if query[1] == 'AND':
            ans = andQuery(query[0], query[2], matrix, words)
            if type(ans) == 'str':
                ans = ans.zfill(len(matrix[0]))
            if ans == -1:
                print(f'\n{query[0]} Not Found.')
            elif ans == -2:
                print(f'\n{query[2]} Not Found.')
            elif ans == '0'*len(matrix[0]):
                print(f'\n{query[0]} AND {query[2]} Not Available.')
            else:
                print(f'\n{query[0]} AND {query[2]} are Available in :')
                for i in range(len(ans)):
                    if ans[i] == '1':
                        print(files[i])

        elif query[1] == 'OR':
                ans = orQuery(query[0], query[2], matrix, words)
                ans = ans.zfill(len(matrix[0]))
                if ans != '0'*len(matrix[0]):
                    print(f'\n{query[0]} OR {query[2]} are Available in :')
                    for i in range(len(ans)):
                        if ans[i] == '1':
                            print(files[i])                
        else:
            print("Invalid Query")
            
    elif len(query) == 2 and query[0] == 'NOT':
        ans = notQuery(query[1], matrix, words)
        if ans != '0'*len(matrix[0]):
            print(f'\n{query[1]}Available in ALL Docs.')
        else:
            print(f'\n{query[1]} is NOT Available in :')
            for i in range(len(ans)):
                if ans[i] == '1':
                    print(files[i])
    else:
        print("Invalid Query")

# 5.
# Implement inverted index for boolean retrieval.
import os
def getFiles():
    lst = []
    for f in os.listdir():      # gives list of files in current directory
        if f.endswith('.txt'):  # taking only .txt files from current directory
            lst.append(f)
    return lst
# create inverted index
def index(files):
    fp = []     # index of this list work as documentID
    for file in files:
        f = open(file, 'r')
        fp.append(f)
    dix = {}    # dictionary: key= terms & value= postings
    for i, f in enumerate(fp):
        dat = f.read().split()
        for word in dat:
            try:
                flag = True # doc not listed for the term
                for j in range(len(dix[word])):
                    if dix[word][j] == i:
                        flag = False # doc already listed for the term
                        break
                    elif dix[word][j] > i:
                        dix[word].insert(j, i)
                        flag = False
                        break
                if flag:
                    dix[word].append(i)
            except KeyError:
                dix[word] = [i] # inserting new term
    return dix
# perform AND Query
def andQuery(dix, w1, w2):
    try:    # get posting for word-1, else error
        l1 = dix[w1]
    except KeyError:
        return -1
    try:    # get posting for word-2, else error
        l2 = dix[w2]
    except KeyError:
        return -2
    # intersection of the two postings
    i, j = 0, 0
    lst = []
    while i < len(l1) and j < len(l2):
        if l1[i] == l2[j]:
            lst.append(l1[i])
            i += 1
            j += 1
        elif l1[i] < l2[j]:
            i += 1
        else:
            j += 1
    return lst
# perform OR Query
def orQuery(dix, w1, w2):
    try:    # get posting for word-1, else works with word-2 only
        l1 = dix[w1]
    except KeyError:
        print(f'{w1} Not Found')
        l1 = []
    try:    # get posting for word-2, else works with word-1 only
        l2 = dix[w2]
    except KeyError:
        print(f'{w2} Not Found')
        l2 = []
    # union of the two postings
    i, j = 0, 0
    lst = []
    while i < len(l1) and j < len(l2):
        if l1[i] == l2[j]:
            lst.append(l1[i])
            i += 1
            j += 1
        elif l1[i] < l2[j]:
            lst.append(l1[i])
            i += 1
        else:
            lst.append(l2[j])
            j += 1
    while i < len(l1):
        lst.append(l1[i])
        i += 1
    while j < len(l2):
        lst.append(l2[j])
        j += 1
    return lst
# perform NOT Query
def notQuery(dix, w, N): 
    try:    # get posting for word
        l = dix[w]
    except KeyError:
        l = []
    lst = [i for i in range(N) if i not in l]
    return lst
if __name__ == '__main__':
    files = getFiles()
    idx = index(files)
    print('\nInverted Index:')
    for i in idx:
        print(i, ':', idx[i])
    print('\nVocabulary')
    for i in idx:
        print(i, end=', ')
    print()

    print('\nTerms in lowercase & Operator in uppercase')
    query = input('Enter query: ').split()
    if len(query) == 3:
        if query[1] == 'AND':
            ans = andQuery(idx, query[0], query[2])
            if ans == -1:
                print(f'{query[0]} Not Found.')
            elif ans == -2:
                print(f'{query[2]} Not Found.')
            elif not ans:
                print(f'{query[0]} AND {query[2]} Not Available.')
            else:
                print(f'{query[0]} AND {query[2]} are Available in :')
                for i in ans:
                    print(files[i])

        elif query[1] == 'OR':
            ans = orQuery(idx, query[0], query[2])
            if ans:
                print(f'{query[0]} OR {query[2]} are Available in :')
                for i in ans:
                    print(files[i])
        else:
            print("Invalid Query")
                        
    elif len(query) == 2 and query[0] == 'NOT':
        ans = notQuery(idx, query[1], len(files))
        if ans:
            print(f'{query[1]} is NOT Available in :')
            for i in ans:
                print(files[i])
        else:
            print(f'\n{query[1]}Available in ALL Docs.')
    else:
        print("Invalid Query")

# 6.
# Using NLTK perform Tokenization, Normalization, Stemming & Lemmetization.
import nltk
nltk.download('punkt')      # download resources for tokenization & puntuations
nltk.download('stopwords')  # download resources for stopwords
nltk.download('wordnet')    # download resources for lemmetization
from nltk.tokenize import sent_tokenize
from nltk.tokenize import word_tokenize
from nltk.probability import FreqDist
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.stem.wordnet import WordNetLemmatizer
from truecase import get_true_case
import string
import pandas as pd
# document
text = "Dr. John Doe visited Paris last week for a conference on artificial intelligence."
print('\nDocument:\n', text)
# Sentence Tokenization
tokenized_sent = sent_tokenize(text)
print('\nSentence Tokenization:\n', tokenized_sent)
# Word Tokenization
tokenized_word = word_tokenize(text)
print('\nWord Tokenization\n', tokenized_word)
# Frequency Distribution
fdist = FreqDist(tokenized_word)
print('\nMost frequent 5 words:\n', fdist.most_common(5))
# Lowercasing
lower_token = []
for token in tokenized_word:
    lower_token.append(token.lower())
print('\nLowercasing:\n', lower_token)
# Truecasing Scentences
true_text = []
for text in tokenized_sent:
    true_text.append(get_true_case(text))
print('\nTruecasing Scentences:\n', true_text)
# Tokenize Truecase Words
true_token = []
for token in true_text:
    true_token.extend(word_tokenize(token))
print('\nTruecase Words\n', true_token)
# Display Original, Lowercase & Truecase Words
d = {'Original': tokenized_word, 'Lowercase': lower_token, 'Truecase': true_token}
df = pd.DataFrame(d)
print('\nOriginal vs Lowercase vs Truecase:\n')
print(df)
# Removing Punctuations
punctuations = list(string.punctuation)
tokens = []
for i in lower_token:
    if i not in punctuations:
        tokens.append(i)
print('\nAfter removing Punctuations:\n', tokens)
# Removing Stopwords
stopwords_english = stopwords.words("english")
filtered_tokens=[]
for w in tokens:
    if w not in stopwords_english:
         filtered_tokens.append(w)
print('\nAfter removing Stopwords:\n', filtered_tokens)
# Stemming
ps = PorterStemmer()
stem_words=[]
for w in filtered_tokens:
     stem_words.append(ps.stem(w))
print('\nAfter Stemming:\n', stem_words)
# Lemmetization
lem = WordNetLemmatizer()
lem_words=[]
for w in filtered_tokens:
     lem_words.append(lem.lemmatize(w))
print('\nAfter Lemmetization:\n', lem_words)
# Display Original, Stemmed & Lemmetized Words
d = {'Original': filtered_tokens, 'Stemming': stem_words, 'Lemmetization': lem_words}
df = pd.DataFrame(d)
print('\nOriginal vs Stemming vs Lemmetization:\n')
print(df)

# 7.
# Naive Bayes classification using scikit-learn.
from sklearn.naive_bayes import GaussianNB, MultinomialNB, BernoulliNB
from sklearn.feature_extraction.text import TfidfVectorizer
import pandas as pd
# documents
doc1 = 'Chinese Beijing Chinese'
doc2 = 'Chinese Chinese Shanghai'
doc3 = 'Chinese Macao'
doc4 = 'Tokyo Japan Chinese'
docs = [doc1, doc2, doc3, doc4]    # doc list
y = [1, 1, 1, 0]                # labels: 1=> chinese, 0=> not chinese
# displaying the training dataset
d = {'Documents': docs, 'Labels': y}
df = pd.DataFrame(d)
print('1=> chinese, 0=> not chinese')
print('Training Dataset:\n', df)
# converting text to numeric using tf-idf
tfidf = TfidfVectorizer()
X = tfidf.fit_transform(docs).toarray()
print('After tf-idf :\n', X)
print('Features:\n', tfidf.get_feature_names_out())
# training Gaussian Naive Bayes
gnb = GaussianNB()
gnb.fit(X, y)
# training Multinomial Naive Bayes
mnb = MultinomialNB()
mnb.fit(X, y)
# training Bernoulli Naive Bayes
bnb = BernoulliNB()
bnb.fit(X, y)
# classifying new document
doc5 = ['Chinese Chinese Chinese Tokyo Japan']
a = tfidf.transform(doc5).toarray()
print('New document: ', doc5[0])
print('Gaussian NB: ', gnb.predict(a)[0])
print('Multinomial NB: ', mnb.predict(a)[0])
print('Bernoulli NB: ', bnb.predict(a)[0])

# 8.
# Rocchio classification using scikit-learn.
import numpy as np
import pandas as pd
from sklearn.neighbors import NearestCentroid
from sklearn.feature_extraction.text import TfidfVectorizer
# documents
doc1 = 'Chinese Beijing Chinese'
doc2 = 'Chinese Chinese Shanghai'
doc3 = 'Chinese Macao'
doc4 = 'Tokyo Japan Chinese'
docs = [doc1, doc2, doc3, doc4]    # doc list
y = [1, 1, 1, 0]                # labels: 1=> chinese, 0=> not chinese
# displaying the training dataset
d = {'Documents': docs, 'Labels': y}
df = pd.DataFrame(d)
print('1=> chinese, 0=> not chinese')
print('Training Dataset:\n', df)
# converting text to numeric using tf-idf
tfidf = TfidfVectorizer()
X = tfidf.fit_transform(docs).toarray()
print('After tf-idf :\n', X)
print('Features:\n', tfidf.get_feature_names_out())
# training rocchio classifier
clf = NearestCentroid()
clf.fit(X, y)
# classifying new document
doc5 = ['Chinese Chinese Chinese Tokyo Japan']
a = tfidf.transform(doc5).toarray()
print('New document: ', doc5[0])
print(clf.predict(a)[0])

# 9.
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
document = [
    "Artificial intelligence is transforming the tech industry.",
    "The latest advancements in machine learning are impressive.",
    "How do you stay healthy and fit with a busy lifestyle?",
    "Cloud computing provides scalable resources over the internet.",
    "Regular exercise is crucial for maintaining good health.",
    "Blockchain technology offers a new level of security.",
    "Eating a balanced diet is essential for overall well-being.",
    "Cybersecurity threats are becoming more sophisticated.",
    "Getting enough sleep is vital for your health.",
    "The future of quantum computing looks promising.",
    "Meditation and mindfulness can reduce stress.",
    "The Internet of Things is connecting devices worldwide.",
    "Staying hydrated is important for your body to function properly.",
    "5G technology will revolutionize communications.",
    "A good mental health routine includes regular relaxation.",
    "Advances in robotics are leading to smarter machines."
]
label = [
    'technology', 'technology', 'health', 'technology', 'health', 
    'technology', 'health', 'technology', 'health', 'technology', 
    'health', 'technology', 'health', 'technology', 'health', 
    'technology'
]
y = [0, 0, 1, 0, 1, 0, 1, 0, 1, 0, 1, 0, 1, 0, 1, 0]
vectorizer = TfidfVectorizer(stop_words='english')
X = vectorizer.fit_transform(document)
true_k = 2
model = KMeans(n_clusters=true_k, init='k-means++', max_iter=100, n_init=1)
model.fit(X)
order_centroids = model.cluster_centers_.argsort()[:, ::-1]
terms = vectorizer.get_feature_names_out()
for i in range(true_k):
    print("\nCluster-%d (top 10 terms):" % i)
    for ind in order_centroids[i, :10]:
        print('%s ' % terms[ind], end=", ")
    print()
X = vectorizer.transform(["The advances in artificial intelligence are remarkable."])
predicted = model.predict(X)[0]
print("\nPrediction = Cluster-", predicted)
# 10.
from nltk import word_tokenize, pos_tag
import nltk
nltk.download('averaged_perceptron_tagger')
sent = "Marie Curie conducted pioneering research on radioactivity in the early 20th century."
tokens = word_tokenize(sent)
print(tokens)
print(pos_tag(tokens))

# PYQ 1. TDM and Cosine Similarity
import nltk
import numpy as np
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Ensure required NLTK data files are downloaded
# nltk.download('punkt')
# nltk.download('stopwords')
# nltk.download('wordnet')
# nltk.download('averaged_perceptron_tagger')

# Sample documents
documents = [
    "Varanasi is a city of masis, culture and religion, is situated in North of India and is one of the oldest living cities of India.",
    "New Delhi is the capital of India and is a modern city. It is the seat of Indian government. It is about 800 kilometers far from Varanasi.",
    "I love Football. This wonderful game is played in countries of the world. I just love watching Football World Cup.",
    "Hockey is the national game of India. It is also played in several countries in the world. It is however losing its popularity to several other games."
]

# Initialize CountVectorizer
vectorizer = CountVectorizer(stop_words='english')
X = vectorizer.fit_transform(documents)

# Convert to dense array for easy manipulation
tdm = X.toarray()
print("Term-Document Matrix:")
print(tdm)

# Function to compute cosine similarity
def compute_cosine_similarity(tdm):
    cos_sim = cosine_similarity(tdm)
    return cos_sim

# Compute cosine similarity
cos_sim = compute_cosine_similarity(tdm)
print("\nCosine Similarity Matrix:")
print(cos_sim)

# Specific cosine similarities
cos_sim_1_2 = cosine_similarity(tdm[0].reshape(1, -1), tdm[1].reshape(1, -1))[0][0]
cos_sim_1_3 = cosine_similarity(tdm[0].reshape(1, -1), tdm[2].reshape(1, -1))[0][0]
cos_sim_1_4 = cosine_similarity(tdm[0].reshape(1, -1), tdm[3].reshape(1, -1))[0][0]
cos_sim_2_4 = cosine_similarity(tdm[1].reshape(1, -1), tdm[3].reshape(1, -1))[0][0]

print(f"\nCosine Similarity between D1 and D2: {cos_sim_1_2}")
print(f"Cosine Similarity between D1 and D3: {cos_sim_1_3}")
print(f"Cosine Similarity between D1 and D4: {cos_sim_1_4}")
print(f"Cosine Similarity between D2 and D4: {cos_sim_2_4}")


# PYQ 2. N-Bayes Classifier

import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import GaussianNB,MultinomialNB,BernoulliNB

doc1 = "Taipei Taiwan"
doc2 = "Macao Taiwan Shanghai"
doc3 = "Japan Sapporo"
doc4 = "Sapporo Osaka Taiwan"

docs=[doc1,doc2,doc3,doc4]
labels=[0,0,1,1]
# 0->china 1->japan


d = {'Documents':docs, "Labels":labels}
df = pd.DataFrame(d)

print(df)

tfidf = TfidfVectorizer()

X = tfidf.fit_transform(docs).toarray()
df2=pd.DataFrame(X, index=docs,columns=tfidf.get_feature_names_out())
# print(df2)

gnb = GaussianNB()
gnb.fit(X,labels)

mnb = MultinomialNB()
mnb.fit(X,labels)


bnb = BernoulliNB()
bnb.fit(X,labels)


doc5 = ["Taiwan Taiwan Sapporo"]

a = tfidf.transform(doc5).toarray()
print(doc5, "\n")
print('\nNew document: ', doc5[0])
print('\nGaussian NB: ', gnb.predict(a)[0], " and probability is ", gnb.predict_proba(a)[0])
print('Multinomial NB: ', mnb.predict(a)[0], " and probability is ", mnb.predict_proba(a)[0])
print('Bernoulli NB: ', bnb.predict(a)[0], " and probability is ", bnb.predict_proba(a)[0])

print("Gaussian NB:")
print(f"  Probability of belonging to class 'China': {gnb.predict_proba(a)[0][0]:.4f}")
print(f"  Probability of belonging to class 'Japan': {gnb.predict_proba(a)[0][1]:.4f}\n")

print("Multinomial NB:")
print(f"  Probability of belonging to class 'China': {mnb.predict_proba(a)[0][0]:.4f}")
print(f"  Probability of belonging to class 'Japan': {mnb.predict_proba(a)[0][1]:.4f}\n")

print("Bernoulli NB:")
print(f"  Probability of belonging to class 'China': {bnb.predict_proba(a)[0][0]:.4f}")
print(f"  Probability of belonging to class 'Japan': {bnb.predict_proba(a)[0][1]:.4f}")


 # PYQ 3. KMeans clustering and Purity 
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from collections import Counter

# Documents and true labels
documents = [
    "This is the most beautiful place in the world.",
    "This man has more skills to show in cricket than any other game.",
    "Hi there! how was your ladakh trip last month?",
    "There was a player who had scored 200+ runs in single cricket innings in his career.",
    "I have got the opportunity to travel to Paris next year for my internship.",
    "Maybe he is better than you in batting but you are much better than him in bowling.",
    "That was really a great day for me when I was there at Lavasa for the whole night.",
    "That's exactly I wanted to become, a highest rating batsmen ever with top scores.",
    "Does it really matter whether you go to Thailand or Goa, it's just you have to spend your holidays.",
    "Why don't you go to Switzerland next year for your 25th Wedding anniversary?",
    "Travel is fatal to prejudice, bigotry, and narrow-mindedness, and many of our people need it sorely on these accounts.",
    "Stop worrying about the potholes in the road and enjoy the journey.",
    "No cricket team in the world depends on one or two players. The team always plays to win.",
    "Cricket is a team game. If you want fame for yourself, go play an individual game.",
    "Because in the end, you won't remember the time you spent working in the office or mowing your lawn. Climb that goddamn mountain.",
    "Isn't cricket supposed to be a team sport? I feel people should decide first whether cricket is a team game or an individual sport."
]
labels = ['travel', 'cricket', 'travel', 'cricket', 'travel', 'cricket', 'travel', 'cricket', 'travel', 'travel', 'travel', 'travel', 'cricket', 'cricket', 'travel', 'cricket']
y = [1, 0, 1, 0, 1, 0, 1, 0, 1, 1, 1, 1, 0, 0, 1, 0]  # 1->travel, 0->cricket

# Vectorize documents using TF-IDF
vectorizer = TfidfVectorizer(stop_words='english')
X = vectorizer.fit_transform(documents)

# Perform KMeans clustering
true_k = 2
model = KMeans(n_clusters=true_k, init='k-means++', max_iter=100, n_init=1)
model.fit(X)

# Print top terms per cluster
order_centroids = model.cluster_centers_.argsort()[:, ::-1]
terms = vectorizer.get_feature_names_out()
for i in range(true_k):
    print("\nCluster-%d (top 10 terms):" % i)
    for ind in order_centroids[i, :10]:
        print('%s ' % terms[ind], end=", ")
    print()

# Predict clusters for each document
clusters = model.predict(X)

# Create a DataFrame with the documents, their true labels, and their predicted clusters
df = pd.DataFrame({'Document': documents, 'TrueLabel': labels, 'Cluster_Predicted': clusters, 'Cluster_True': y})
print("\nDocument clustering results:\n")
print(df)

# Calculate purity
def purity_score(y_true, y_pred):
    # Create a confusion matrix
    contingency_matrix = pd.crosstab(y_pred, y_true)
    # Return purity
    return np.sum(np.amax(contingency_matrix.values, axis=1)) / np.sum(contingency_matrix.values)

purity = purity_score(y, clusters)
print(f"\nPurity: {purity:.4f}")








